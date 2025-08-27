import os
import sqlite3
import psycopg2
from flask import Flask, render_template, request, redirect, session, send_file, g
from docx import Document
from io import BytesIO
from datetime import date, timedelta

app = Flask(__name__)
# Best practice: use environment variable for secret key, with a local fallback
app.secret_key = os.environ.get("SECRET_KEY", "your-super-secret-key-here")

# ================== Database Connection ==================
# This function connects to the appropriate database.
# It checks for the DATABASE_URL environment variable (from Render).
# If found, it connects to PostgreSQL (Supabase).
# If not, it falls back to a local SQLite database for offline use.
def get_db():
    # Use g object to store the connection and avoid multiple connections per request
    if 'db' not in g:
        DATABASE_URL = os.environ.get('DATABASE_URL')
        if DATABASE_URL:
            # PostgreSQL connection for production (Render)
            # The psycopg2 library needs the correct connection string format
            if DATABASE_URL.startswith("postgresql://"):
                DATABASE_URL = DATABASE_URL.replace("postgresql://", "postgres://", 1)
            try:
                g.db = psycopg2.connect(DATABASE_URL)
                g.db_type = 'postgresql'
            except Exception as e:
                print(f"Error connecting to PostgreSQL database: {e}")
                # Fallback to SQLite if connection fails
                g.db = sqlite3.connect('database.db')
                g.db_type = 'sqlite'
                print("Falling back to local SQLite database.")
        else:
            # SQLite connection for local development
            g.db = sqlite3.connect('database.db')
            g.db_type = 'sqlite'
            # Enable foreign key support for SQLite
            g.db.execute('PRAGMA foreign_keys = ON')

    return g.db

# Closes the database connection at the end of the request
@app.teardown_appcontext
def close_connection(exception):
    db = g.pop('db', None)
    if db is not None:
        db.close()

# This function initializes the database tables if they don't exist.
def init_db():
    conn = get_db()
    if conn is None:
        print("Could not connect to the database to initialize it.")
        return

    cur = conn.cursor()
    
    # Use different SQL syntax based on the database type
    if g.db_type == 'sqlite':
        cur.execute("""
            CREATE TABLE IF NOT EXISTS usuarios (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                usuario TEXT UNIQUE,
                nombre TEXT,
                apellido TEXT,
                rol TEXT,
                clave TEXT,
                perfil TEXT
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS cursos (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                nombre TEXT,
                año INTEGER
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS alumnos (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                nombre TEXT,
                apellido TEXT,
                curso_id INTEGER,
                FOREIGN KEY(curso_id) REFERENCES cursos(id)
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS docente_cursos (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                docente_id INTEGER,
                curso_id INTEGER,
                FOREIGN KEY(docente_id) REFERENCES usuarios(id),
                FOREIGN KEY(curso_id) REFERENCES cursos(id)
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS notas (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                alumno_id INTEGER,
                docente_id INTEGER,
                curso_id INTEGER,
                nota REAL,
                fecha TEXT,
                FOREIGN KEY(alumno_id) REFERENCES alumnos(id)
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS asistencia (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                alumno_id INTEGER,
                docente_id INTEGER,
                curso_id INTEGER,
                fecha TEXT,
                presente INTEGER,
                FOREIGN KEY(alumno_id) REFERENCES alumnos(id)
            )
        """)
        
        # Create the default admin user if it doesn't exist
        cur.execute("SELECT * FROM usuarios WHERE rol='admin'")
        if not cur.fetchone():
            cur.execute("""
                INSERT INTO usuarios (usuario, nombre, apellido, rol, clave, perfil)
                VALUES (?, ?, ?, ?, ?, ?)
            """, ("admin", "Admin", "Taller", "admin", "1234", ""))
            print("Admin user created: user=admin, password=1234")

    elif g.db_type == 'postgresql':
        cur.execute("""
            CREATE TABLE IF NOT EXISTS usuarios (
                id SERIAL PRIMARY KEY,
                usuario TEXT UNIQUE,
                nombre TEXT,
                apellido TEXT TEXT,
                rol TEXT,
                clave TEXT,
                perfil TEXT
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS cursos (
                id SERIAL PRIMARY KEY,
                nombre TEXT,
                año INTEGER
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS alumnos (
                id SERIAL PRIMARY KEY,
                nombre TEXT,
                apellido TEXT,
                curso_id INTEGER,
                FOREIGN KEY(curso_id) REFERENCES cursos(id)
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS docente_cursos (
                id SERIAL PRIMARY KEY,
                docente_id INTEGER,
                curso_id INTEGER,
                FOREIGN KEY(docente_id) REFERENCES usuarios(id),
                FOREIGN KEY(curso_id) REFERENCES cursos(id)
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS notas (
                id SERIAL PRIMARY KEY,
                alumno_id INTEGER,
                docente_id INTEGER,
                curso_id INTEGER,
                nota REAL,
                fecha TEXT,
                FOREIGN KEY(alumno_id) REFERENCES alumnos(id)
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS asistencia (
                id SERIAL PRIMARY KEY,
                alumno_id INTEGER,
                docente_id INTEGER,
                curso_id INTEGER,
                fecha TEXT,
                presente INTEGER,
                FOREIGN KEY(alumno_id) REFERENCES alumnos(id)
            )
        """)
        
        # Create the default admin user if it doesn't exist
        cur.execute("SELECT * FROM usuarios WHERE rol='admin'")
        if not cur.fetchone():
            cur.execute("""
                INSERT INTO usuarios (usuario, nombre, apellido, rol, clave, perfil)
                VALUES (%s, %s, %s, %s, %s, %s)
            """, ("admin", "Admin", "Taller", "admin", "1234", ""))
            print("Admin user created: user=admin, password=1234")

    conn.commit()
    cur.close()


with app.app_context():
    init_db()


# ================== Login ==================
@app.route("/", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        usuario_form = request.form["usuario"]
        clave = request.form["clave"]
        conn = get_db()
        cur = conn.cursor()
        if g.db_type == 'sqlite':
            cur.execute("SELECT * FROM usuarios WHERE usuario=? AND clave=?", (usuario_form, clave))
        else:
            cur.execute("SELECT * FROM usuarios WHERE usuario=%s AND clave=%s", (usuario_form, clave))
        
        usuario = cur.fetchone()
        cur.close()
        conn.close()

        if usuario:
            session["usuario_id"] = usuario[0]
            session["rol"] = usuario[4]
            if usuario[4] == "admin":
                return redirect("/admin")
            else:
                return redirect("/docente")
        else:
            return "Usuario o clave incorrecta"
    return render_template("login.html")


# ================== Admin ==================
@app.route("/admin")
def admin():
    if "rol" in session and session["rol"] == "admin":
        conn = get_db()
        cur = conn.cursor()

        cur.execute("SELECT * FROM cursos")
        cursos = cur.fetchall()

        if g.db_type == 'sqlite':
            cur.execute("""
                SELECT u.id, u.nombre, u.apellido, u.usuario, c.nombre, c.año
                FROM usuarios u
                LEFT JOIN docente_cursos dc ON u.id = dc.docente_id
                LEFT JOIN cursos c ON dc.curso_id = c.id
                WHERE u.rol='docente'
            """)
        else:
            cur.execute("""
                SELECT u.id, u.nombre, u.apellido, u.usuario, c.nombre, c.año
                FROM usuarios u
                LEFT JOIN docente_cursos dc ON u.id = dc.docente_id
                LEFT JOIN cursos c ON dc.curso_id = c.id
                WHERE u.rol='docente'
            """)
        docentes = cur.fetchall()

        if g.db_type == 'sqlite':
            cur.execute("""
                SELECT c.id, c.nombre, c.año, a.id, a.nombre, a.apellido
                FROM cursos c
                LEFT JOIN alumnos a ON a.curso_id = c.id
                ORDER BY c.id, a.apellido, a.nombre
            """)
        else:
            cur.execute("""
                SELECT c.id, c.nombre, c.año, a.id, a.nombre, a.apellido
                FROM cursos c
                LEFT JOIN alumnos a ON a.curso_id = c.id
                ORDER BY c.id, a.apellido, a.nombre
            """)
        alumnos_curso_raw = cur.fetchall()

        cur.close()

        alumnos_curso = {}
        for curso_id, curso_nombre, curso_año, alumno_id, alumno_nombre, alumno_apellido in alumnos_curso_raw:
            if curso_id not in alumnos_curso:
                alumnos_curso[curso_id] = {
                    "nombre": curso_nombre,
                    "año": curso_año,
                    "alumnos": []
                }
            if alumno_id:
                alumnos_curso[curso_id]["alumnos"].append({
                    "id": alumno_id,
                    "nombre": alumno_nombre,
                    "apellido": alumno_apellido
                })

        hoy = date.today()
        lunes_actual = hoy - timedelta(days=hoy.weekday())
        fecha_inicio_default = lunes_actual.isoformat()

        return render_template("admin.html",
                               cursos=cursos,
                               docentes=docentes,
                               alumnos_curso=alumnos_curso,
                               fecha_inicio_default=fecha_inicio_default)
    return redirect("/")


# ================== Agregar curso ==================
@app.route("/agregar_curso", methods=["GET", "POST"])
def agregar_curso():
    if "rol" in session and session["rol"] == "admin":
        if request.method == "POST":
            nombre = request.form["nombre"]
            año = request.form["año"]
            conn = get_db()
            cur = conn.cursor()
            if g.db_type == 'sqlite':
                cur.execute("INSERT INTO cursos (nombre, año) VALUES (?,?)", (nombre, año))
            else:
                cur.execute("INSERT INTO cursos (nombre, año) VALUES (%s,%s)", (nombre, año))
            conn.commit()
            cur.close()
            return redirect("/admin")
        return render_template("agregar_curso.html")
    return redirect("/")


# ================== Agregar docente ==================
@app.route("/agregar_docente", methods=["GET", "POST"])
def agregar_docente():
    if "rol" in session and session["rol"] == "admin":
        conn = get_db()
        cur = conn.cursor()
        cur.execute("SELECT * FROM cursos")
        cursos = cur.fetchall()

        if request.method == "POST":
            usuario = request.form["usuario"].strip()
            nombre = request.form["nombre"].strip()
            apellido = request.form["apellido"].strip()
            clave = request.form["clave"].strip()
            perfil = request.form["perfil"].strip()
            curso_id = request.form["curso"]

            if g.db_type == 'sqlite':
                cur.execute("SELECT id FROM usuarios WHERE usuario=?", (usuario,))
            else:
                cur.execute("SELECT id FROM usuarios WHERE usuario=%s", (usuario,))
            usuario_existente = cur.fetchone()

            if usuario_existente:
                docente_id = usuario_existente[0]
            else:
                if g.db_type == 'sqlite':
                    cur.execute(
                        "INSERT INTO usuarios (usuario, nombre, apellido, rol, clave, perfil) VALUES (?,?,?,?,?,?)",
                        (usuario, nombre, apellido, "docente", clave, perfil)
                    )
                    docente_id = cur.lastrowid
                else:
                    cur.execute(
                        "INSERT INTO usuarios (usuario, nombre, apellido, rol, clave, perfil) VALUES (%s,%s,%s,%s,%s,%s) RETURNING id",
                        (usuario, nombre, apellido, "docente", clave, perfil)
                    )
                    docente_id = cur.fetchone()[0]

            if g.db_type == 'sqlite':
                cur.execute("SELECT * FROM docente_cursos WHERE docente_id=? AND curso_id=?", (docente_id, curso_id))
            else:
                cur.execute("SELECT * FROM docente_cursos WHERE docente_id=%s AND curso_id=%s", (docente_id, curso_id))
            if not cur.fetchone():
                if g.db_type == 'sqlite':
                    cur.execute("INSERT INTO docente_cursos (docente_id, curso_id) VALUES (?,?)", (docente_id, curso_id))
                else:
                    cur.execute("INSERT INTO docente_cursos (docente_id, curso_id) VALUES (%s,%s)", (docente_id, curso_id))

            conn.commit()
            cur.close()
            return redirect("/admin")

        cur.close()
        return render_template("agregar_docente.html", cursos=cursos)
    return redirect("/")


# ================== Agregar alumno ==================
@app.route("/agregar_alumno", methods=["GET", "POST"])
def agregar_alumno():
    if "rol" in session and session["rol"] == "admin":
        conn = get_db()
        cur = conn.cursor()
        cur.execute("SELECT * FROM cursos")
        cursos = cur.fetchall()
        if request.method == "POST":
            nombre = request.form["nombre"]
            apellido = request.form["apellido"]
            curso_id = request.form["curso"]
            if g.db_type == 'sqlite':
                cur.execute("INSERT INTO alumnos (nombre, apellido, curso_id) VALUES (?,?,?)",
                            (nombre, apellido, curso_id))
            else:
                cur.execute("INSERT INTO alumnos (nombre, apellido, curso_id) VALUES (%s,%s,%s)",
                            (nombre, apellido, curso_id))
            conn.commit()
            cur.close()
            return redirect("/admin")
        cur.close()
        return render_template("agregar_alumno.html", cursos=cursos)
    return redirect("/")


# ================== Docente ==================
@app.route("/docente")
def docente():
    if "rol" in session and session["rol"] == "docente":
        docente_id = session["usuario_id"]
        conn = get_db()
        cur = conn.cursor()
        if g.db_type == 'sqlite':
            cur.execute("""
                SELECT dc.id, c.nombre, c.año, c.id
                FROM docente_cursos dc
                JOIN cursos c ON c.id = dc.curso_id
                WHERE dc.docente_id=?
            """, (docente_id,))
        else:
            cur.execute("""
                SELECT dc.id, c.nombre, c.año, c.id
                FROM docente_cursos dc
                JOIN cursos c ON c.id = dc.curso_id
                WHERE dc.docente_id=%s
            """, (docente_id,))
        asignaciones_raw = cur.fetchall()
        cur.close()

        asignaciones = [(row[3], row[1], row[2]) for row in asignaciones_raw]

        hoy = date.today().isoformat()

        return render_template("docente.html", asignaciones=asignaciones, today=hoy)
    return redirect("/")


# ================== Notas ==================
@app.route("/notas/<int:curso_id>", methods=["GET", "POST"])
def notas(curso_id):
    if "rol" in session and session["rol"] == "docente":
        docente_id = session["usuario_id"]
        conn = get_db()
        cur = conn.cursor()
        if g.db_type == 'sqlite':
            cur.execute("SELECT * FROM alumnos WHERE curso_id=?", (curso_id,))
        else:
            cur.execute("SELECT * FROM alumnos WHERE curso_id=%s", (curso_id,))
        alumnos = cur.fetchall()

        if request.method == "POST":
            fecha = date.today().isoformat()
            for alumno in alumnos:
                nota = request.form.get(f"nota_{alumno[0]}")
                if nota:
                    if g.db_type == 'sqlite':
                        cur.execute(
                            "INSERT INTO notas (alumno_id, docente_id, curso_id, nota, fecha) VALUES (?,?,?,?,?)",
                            (alumno[0], docente_id, curso_id, float(nota), fecha)
                        )
                    else:
                        cur.execute(
                            "INSERT INTO notas (alumno_id, docente_id, curso_id, nota, fecha) VALUES (%s,%s,%s,%s,%s)",
                            (alumno[0], docente_id, curso_id, float(nota), fecha)
                        )
            conn.commit()
            cur.close()
            return "Notas registradas correctamente"

        cur.close()
        return render_template("notas.html", alumnos=alumnos, curso_id=curso_id)
    return redirect("/")


# ================== Asistencia ==================
@app.route("/asistencia/<int:curso_id>", methods=["GET", "POST"])
def asistencia(curso_id):
    if "rol" not in session or session["rol"] != "docente":
        return redirect("/")

    docente_id = session["usuario_id"]

    inicio_semana_str = request.args.get("inicio")
    if inicio_semana_str:
        inicio_semana = date.fromisoformat(inicio_semana_str)
    else:
        hoy = date.today()
        inicio_semana = hoy - timedelta(days=hoy.weekday())

    fechas_semana = [inicio_semana + timedelta(days=i) for i in range(5)]

    conn = get_db()
    cur = conn.cursor()

    if g.db_type == 'sqlite':
        cur.execute("SELECT id, nombre, apellido FROM alumnos WHERE curso_id=?", (curso_id,))
    else:
        cur.execute("SELECT id, nombre, apellido FROM alumnos WHERE curso_id=%s", (curso_id,))
    alumnos = cur.fetchall()

    if request.method == "POST":
        for alumno in alumnos:
            for f in fechas_semana:
                presente = request.form.get(f"asistencia_{alumno[0]}_{f}")
                if g.db_type == 'sqlite':
                    cur.execute("""
                        SELECT id FROM asistencia
                        WHERE alumno_id=? AND docente_id=? AND curso_id=? AND fecha=?
                    """, (alumno[0], docente_id, curso_id, f.isoformat()))
                else:
                    cur.execute("""
                        SELECT id FROM asistencia
                        WHERE alumno_id=%s AND docente_id=%s AND curso_id=%s AND fecha=%s
                    """, (alumno[0], docente_id, curso_id, f.isoformat()))
                existing = cur.fetchone()
                if existing:
                    if g.db_type == 'sqlite':
                        cur.execute("UPDATE asistencia SET presente=? WHERE id=?", (1 if presente else 0, existing[0]))
                    else:
                        cur.execute("UPDATE asistencia SET presente=%s WHERE id=%s", (1 if presente else 0, existing[0]))
                else:
                    if g.db_type == 'sqlite':
                        cur.execute("""
                            INSERT INTO asistencia (alumno_id, docente_id, curso_id, fecha, presente)
                            VALUES (?, ?, ?, ?, ?)
                        """, (alumno[0], docente_id, curso_id, f.isoformat(), 1 if presente else 0))
                    else:
                        cur.execute("""
                            INSERT INTO asistencia (alumno_id, docente_id, curso_id, fecha, presente)
                            VALUES (%s, %s, %s, %s, %s)
                        """, (alumno[0], docente_id, curso_id, f.isoformat(), 1 if presente else 0))
        conn.commit()
        cur.close()
        return redirect(f"/asistencia/{curso_id}?inicio={inicio_semana.isoformat()}")

    asistencia = {}
    for alumno in alumnos:
        asistencia[alumno[0]] = {}
        for f in fechas_semana:
            if g.db_type == 'sqlite':
                cur.execute("""
                    SELECT presente FROM asistencia
                    WHERE alumno_id=? AND docente_id=? AND curso_id=? AND fecha=?
                """, (alumno[0], docente_id, curso_id, f.isoformat()))
            else:
                cur.execute("""
                    SELECT presente FROM asistencia
                    WHERE alumno_id=%s AND docente_id=%s AND curso_id=%s AND fecha=%s
                """, (alumno[0], docente_id, curso_id, f.isoformat()))
            res = cur.fetchone()
            asistencia[alumno[0]][f] = res[0] if res else 0

    cur.close()
    dias_semana = ["Lunes", "Martes", "Miércoles", "Jueves", "Viernes"]
    fechas_semana_nombres = [(f, dias_semana[f.weekday()]) for f in fechas_semana]

    semana_anterior = inicio_semana - timedelta(days=7)
    semana_siguiente = inicio_semana + timedelta(days=7)

    return render_template(
        "asistencia.html",
        alumnos=alumnos,
        fechas=fechas_semana_nombres,
        asistencia=asistencia,
        curso_id=curso_id,
        semana_anterior=semana_anterior,
        semana_siguiente=semana_siguiente
    )


# ================== Exportar Notas ==================
@app.route("/exportar_notas/<int:curso_id>")
def exportar_notas(curso_id):
    conn = get_db()
    cur = conn.cursor()

    if g.db_type == 'sqlite':
        cur.execute("SELECT nombre, año FROM cursos WHERE id=?", (curso_id,))
    else:
        cur.execute("SELECT nombre, año FROM cursos WHERE id=%s", (curso_id,))
    curso = cur.fetchone()
    curso_nombre = f"{curso[0]} - Año {curso[1]}" if curso else "Curso Desconocido"

    doc = Document()
    doc.add_heading(f"Notas del Curso: {curso_nombre}", 0)

    if g.db_type == 'sqlite':
        cur.execute("""
            SELECT a.apellido || ', ' || a.nombre, n.nota
            FROM alumnos a
            LEFT JOIN notas n ON a.id = n.alumno_id
            WHERE a.curso_id=?
            ORDER BY a.apellido, a.nombre
        """, (curso_id,))
    else:
        cur.execute("""
            SELECT a.apellido || ', ' || a.nombre, n.nota
            FROM alumnos a
            LEFT JOIN notas n ON a.id = n.alumno_id
            WHERE a.curso_id=%s
            ORDER BY a.apellido, a.nombre
        """, (curso_id,))
    resultados = cur.fetchall()

    cur.close()

    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Alumno'
    hdr_cells[1].text = 'Nota'

    for alumno, nota in resultados:
        row_cells = table.add_row().cells
        row_cells[0].text = alumno
        row_cells[1].text = str(nota if nota is not None else "")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(buffer, as_attachment=True, download_name=f"Notas_{curso_nombre}.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# ================== Exportar Asistencia ==================
@app.route("/exportar_asistencia/<int:curso_id>")
def exportar_asistencia(curso_id):
    inicio_semana_str = request.args.get("inicio")
    if inicio_semana_str:
        inicio_semana = date.fromisoformat(inicio_semana_str)
    else:
        hoy = date.today()
        inicio_semana = hoy - timedelta(days=hoy.weekday())

    fechas_semana = [inicio_semana + timedelta(days=i) for i in range(5)]
    dias_semana = ["Lunes", "Martes", "Miércoles", "Jueves", "Viernes"]
    fechas_semana_nombres = [(f, dias_semana[f.weekday()]) for f in fechas_semana]

    conn = get_db()
    cur = conn.cursor()

    if g.db_type == 'sqlite':
        cur.execute("SELECT nombre, año FROM cursos WHERE id=?", (curso_id,))
    else:
        cur.execute("SELECT nombre, año FROM cursos WHERE id=%s", (curso_id,))
    curso = cur.fetchone()
    if not curso:
        cur.close()
        return "Curso no encontrado"
    curso_nombre, curso_año = curso

    if g.db_type == 'sqlite':
        cur.execute(
            "SELECT id, apellido, nombre FROM alumnos WHERE curso_id=? ORDER BY apellido, nombre",
            (curso_id,)
        )
    else:
        cur.execute(
            "SELECT id, apellido, nombre FROM alumnos WHERE curso_id=%s ORDER BY apellido, nombre",
            (curso_id,)
        )
    alumnos = cur.fetchall()
    cur.close()

    doc = Document()
    doc.add_heading(f"Asistencia del Curso: {curso_nombre} - Año {curso_año}", 0)
    doc.add_paragraph(
        f"Semana: {inicio_semana.strftime('%d/%m/%Y')} - "
        f"{(inicio_semana + timedelta(days=4)).strftime('%d/%m/%Y')}"
    )
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=1 + len(fechas_semana))
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Alumno"
    for i, (fecha, nombre_dia) in enumerate(fechas_semana_nombres):
        hdr_cells[i + 1].text = f"{nombre_dia}\n{fecha.strftime('%d/%m')}"

    for alumno_id, apellido, nombre in alumnos:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{apellido}, {nombre}"
        for j, (fecha, nombre_dia) in enumerate(fechas_semana_nombres):
            conn = get_db()
            cur = conn.cursor()
            if g.db_type == 'sqlite':
                cur.execute("""
                    SELECT presente FROM asistencia
                    WHERE alumno_id=? AND curso_id=? AND fecha=?
                """, (alumno_id, curso_id, fecha.isoformat()))
            else:
                cur.execute("""
                    SELECT presente FROM asistencia
                    WHERE alumno_id=%s AND curso_id=%s AND fecha=%s
                """, (alumno_id, curso_id, fecha.isoformat()))
            res = cur.fetchone()
            cur.close()
            if res is None:
                estado = "SR"
            else:
                estado = "P" if res[0] == 1 else "A"
            row_cells[j + 1].text = estado

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"Asistencia_{curso_nombre}_{inicio_semana.strftime('%d-%m-%Y')}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# ================== Exportar Alumnos ==================
@app.route("/exportar_alumnos/<int:curso_id>")
def exportar_alumnos(curso_id):
    doc = Document()

    conn = get_db()
    cur = conn.cursor()
    if g.db_type == 'sqlite':
        cur.execute("SELECT nombre, año FROM cursos WHERE id=?", (curso_id,))
    else:
        cur.execute("SELECT nombre, año FROM cursos WHERE id=%s", (curso_id,))
    curso = cur.fetchone()
    if not curso:
        cur.close()
        return "Curso no encontrado"
    curso_nombre, curso_año = curso

    doc.add_heading(f"Alumnos del Curso: {curso_nombre} - Año {curso_año}", 0)

    if g.db_type == 'sqlite':
        cur.execute("""
            SELECT apellido, nombre
            FROM alumnos
            WHERE curso_id=?
            ORDER BY apellido, nombre
        """, (curso_id,))
    else:
        cur.execute("""
            SELECT apellido, nombre
            FROM alumnos
            WHERE curso_id=%s
            ORDER BY apellido, nombre
        """, (curso_id,))
    alumnos = cur.fetchall()

    cur.close()

    for apellido, nombre in alumnos:
        doc.add_paragraph(f"{apellido}, {nombre}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"Alumnos_Curso_{curso_nombre}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# ================== Eliminar Curso ==================
@app.route("/eliminar_curso/<int:curso_id>", methods=["POST"])
def eliminar_curso(curso_id):
    conn = get_db()
    cur = conn.cursor()
    if g.db_type == 'sqlite':
        cur.execute("DELETE FROM asistencia WHERE curso_id=?", (curso_id,))
        cur.execute("DELETE FROM alumnos WHERE curso_id=?", (curso_id,))
        cur.execute("DELETE FROM cursos WHERE id=?", (curso_id,))
    else:
        cur.execute("DELETE FROM asistencia WHERE curso_id=%s", (curso_id,))
        cur.execute("DELETE FROM alumnos WHERE curso_id=%s", (curso_id,))
        cur.execute("DELETE FROM cursos WHERE id=%s", (curso_id,))
    conn.commit()
    cur.close()
    return redirect("/admin")


# ================== Eliminar Docente ==================
@app.route("/eliminar_docente/<int:docente_id>", methods=["POST"])
def eliminar_docente(docente_id):
    conn = get_db()
    cur = conn.cursor()
    if g.db_type == 'sqlite':
        cur.execute("DELETE FROM docente_cursos WHERE docente_id=?", (docente_id,))
        cur.execute("DELETE FROM usuarios WHERE id=? AND rol='docente'", (docente_id,))
    else:
        cur.execute("DELETE FROM docente_cursos WHERE docente_id=%s", (docente_id,))
        cur.execute("DELETE FROM usuarios WHERE id=%s AND rol='docente'", (docente_id,))
    conn.commit()
    cur.close()
    return redirect("/admin")


# ================== Eliminar Alumno ==================
@app.route("/eliminar_alumno/<int:alumno_id>", methods=["POST"])
def eliminar_alumno(alumno_id):
    conn = get_db()
    cur = conn.cursor()
    if g.db_type == 'sqlite':
        cur.execute("DELETE FROM asistencia WHERE alumno_id=?", (alumno_id,))
        cur.execute("DELETE FROM alumnos WHERE id=?", (alumno_id,))
    else:
        cur.execute("DELETE FROM asistencia WHERE alumno_id=%s", (alumno_id,))
        cur.execute("DELETE FROM alumnos WHERE id=%s", (alumno_id,))
    conn.commit()
    cur.close()
    return redirect("/admin")


# ================== Cerrar sesión ==================
@app.route("/logout")
def logout():
    session.clear()
    return redirect("/")


if __name__ == "__main__":
    app.run(debug=True)
